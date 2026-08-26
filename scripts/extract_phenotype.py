#!/usr/bin/env python3
"""Extract plain text from the private phenotype DOCX into runtime/."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", required=True)
    parser.add_argument("--out", default="runtime/phenotype.txt")
    args = parser.parse_args()

    document = Document(args.docx)
    lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                lines.append(" | ".join(values))
    output = Path(args.out)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text("\n".join(lines) + "\n")
    print(f"wrote {len(lines)} lines to {output}")


if __name__ == "__main__":
    main()

